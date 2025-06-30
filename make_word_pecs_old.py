import os
import json
import torch
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Mm
import io
from docx.shared import Inches




# Configuration
RESULTS_ROOT = "long_run_T_0-6"
PLOTS_ROOT = os.path.join(RESULTS_ROOT, "plots")
DOCX_PATH = os.path.join(RESULTS_ROOT, "summary_report.docx")
SELECTED_LAYERS = ["first", "last"]
THRESHOLD = 0.3

# Ensure plots directory exists
os.makedirs(PLOTS_ROOT, exist_ok=True)

# Initialize Word document
doc = Document()
doc.add_heading("Generation Report", level=0)

# Process each prompt folder
for prompt_name in os.listdir(RESULTS_ROOT):
    prompt_dir = os.path.join(RESULTS_ROOT, prompt_name)
    if not os.path.isdir(prompt_dir) or prompt_name == "plots":
        continue

    # Load result.json
    result_file = os.path.join(prompt_dir, "result.json")
    with open(result_file) as f:
        result = json.load(f)
    prompt_text = result['prompt']
    generated_text = result['generated_text']

    # Add section for this prompt
    doc.add_heading(prompt_name.replace('_', ' ').title(), level=1)
    doc.add_heading("Prompt", level=2)
    doc.add_paragraph(prompt_text)
    doc.add_heading("Generated Text", level=2)
    doc.add_paragraph(generated_text)

    # For each pair of selected layers: plot cos sim and l2
    for i, layer in enumerate(SELECTED_LAYERS):
        # Filenames
        cos_name = f"cosine_sim_matrix_{layer}.pt"
        cos_path = os.path.join(prompt_dir, cos_name)

        # Load tensors
        cos_sim = torch.load(cos_path)

        # Plot cosine similarity recurrence
        plt.figure(figsize=(6, 6))
        plt.title(f"Cosine Similarity: {layer}")
        plt.xlabel("Time Index")
        plt.ylabel("Time Index")
        plt.imshow(cos_sim.cpu().numpy(), origin='lower', aspect='equal')
        plt.colorbar()
        cos_plot = os.path.join(PLOTS_ROOT, f"{prompt_name}_cos_{layer}.png")
        # plt.savefig(cos_plot, dpi=300)
        # doc.add_picture(cos_plot, width=Mm(120))


        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png')  # Save to buffer instead of file
        img_stream.seek(0)  # Rewind to the beginning
        doc.add_heading(f"Cosine Similarity ({layer})", level=3)
        doc.add_picture(img_stream, width=Mm(120))

        plt.close()

        # Recurrence plot from cosine similarity
        # Compute normalized distance
        distances = 1 - cos_sim
        distances = distances / distances.max()
        recurrence = distances < THRESHOLD

        plt.figure(figsize=(6, 6))
        plt.title(f"Recurrence Plot {layer}\nThreshold: {THRESHOLD}")
        plt.xlabel("Time Index")
        plt.ylabel("Time Index")
        plt.imshow(recurrence.cpu().numpy(), cmap='binary', origin='lower')
        plt.colorbar(label="Recurrence")
        rec_plot = os.path.join(PLOTS_ROOT, f"{prompt_name}_rec_{layer}.png")
        # plt.savefig(rec_plot, dpi=300)
        # doc.add_picture(rec_plot, width=Mm(120))

        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png')  # Save to buffer instead of file
        img_stream.seek(0)  # Rewind to the beginning
        doc.add_heading(f"Recurrence Plot ({layer}, threshold={THRESHOLD})", level=3)
        doc.add_picture(img_stream, width=Mm(120))

        plt.close()

    doc.add_page_break()

# Save document
doc.save(DOCX_PATH)
print(f"Report generated: {DOCX_PATH}")
