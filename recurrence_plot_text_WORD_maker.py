import matplotlib.pyplot as plt
import torch
import os
import pickle
import json
# import base64
# import io
from docxtpl import DocxTemplate #, InlineImage
from docx.shared import Mm
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet #, ParagraphStyle
from reportlab.lib.units import inch


radiuses = [0.01, 0.02, 0.04]
temps = [0, 0.6]
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

def load_pickle(name):
    with open(os.path.join(RESULTS_DIR, f"{name}.pkl"), "rb") as f:
        return pickle.load(f)

def recurrence_plot_with_threshold(trajectories, output_dir="./", results=None):
    """Generate recurrence plots and save to Word document"""
    
    # First, create a blank template document
    template_path = os.path.join(output_dir, "template.docx")
    
    # Create a simple template document if it doesn't exist
    if not os.path.exists(template_path):
        temp_doc = Document()
        temp_doc.save(template_path)

    doc = DocxTemplate(template_path)
    
    
    # Create a context dictionary for the template
    context = {
        'title': f"Recurrence Plots Analysis - T={TEMPERATURE}, R={RADIUS_INITIAL_CONDITIONS}",
        'trajectories': []
    }
    
    # Add initial prompt if available
    if results and "initial_prompt" in results:
        context['initial_prompt'] = results["initial_prompt"]
    
    # Process each trajectory
    for traj_idx, traj in enumerate(trajectories):
        traj_data = {'index': traj_idx}
        
        # Add generated text for this trajectory if available
        if results and "generated_texts" in results and traj_idx < len(results["generated_texts"]):
            traj_data['generated_text'] = results["generated_texts"][traj_idx]
        
        ### DISTANCE WITH ITSELF (typical recurrence plot)
        # cosine 
        traj_normalized = traj / torch.norm(traj, dim=1, keepdim=True)
        cosine_sim = traj_normalized @ traj_normalized.T  # (n, n)
        distances = 1 - cosine_sim
        distances = distances / torch.max(distances)
        
        # Save distance plot to file and to memory
        plt.figure(figsize=(10, 8))
        plt.title(f"Distance Plot (Trajectory {traj_idx})")
        plt.xlabel("Time Index")
        plt.ylabel("Time Index")
        plt.imshow(distances.cpu().numpy(), cmap='coolwarm', aspect='equal', vmin=-1, vmax=1, origin='lower')
        plt.colorbar()
        
        # Save to file for reference
        distance_plot_path = os.path.join(output_dir, f"recurrence_traj_{traj_idx}_cos_distance.png")
        plt.savefig(distance_plot_path, dpi=300)
        
        # Convert to InlineImage
        traj_data['distance_plot_path'] = distance_plot_path
        plt.close()
        
        # For recurrence plot
        threshold = 0.3
        recurrence = distances < threshold
        
        # Plot recurrence plot
        plt.figure(figsize=(10, 8))
        plt.imshow(recurrence.cpu().numpy(), cmap='binary', origin='lower')
        plt.title(f"Recurrence Plot (Trajectory {traj_idx})\n Threshold: {threshold}")
        plt.xlabel("Time Index")
        plt.ylabel("Time Index")
        plt.colorbar(label="Recurrence")
        
        # Save to file for reference
        recurrence_plot_path = os.path.join(output_dir, f"recurrence_traj_{traj_idx}_cos_threshold_{threshold}.png")
        plt.savefig(recurrence_plot_path, dpi=300)
        
        # Add to trajectory data
        traj_data['recurrence_plot_path'] = recurrence_plot_path
        plt.close()
        
        # Add trajectory data to context
        context['trajectories'].append(traj_data)
    
    # Create the document from scratch since we're having issues with docx
    try:
        # First try using a standard Word creation approach with python-docx
        doc = Document()
        
        # Add title
        doc.add_heading(context['title'], level=0)
        
        # Add initial prompt if available
        if 'initial_prompt' in context:
            doc.add_heading("Initial Prompt", level=1)
            doc.add_paragraph(context['initial_prompt'])
            doc.add_page_break()
        
        # Add trajectories
        for traj in context['trajectories']:
            idx = traj['index']
            doc.add_heading(f"Trajectory {idx}", level=1)
            
            # Add generated text if available
            if 'generated_text' in traj:
                doc.add_heading("Generated Text:", level=2)
                doc.add_paragraph(traj['generated_text'])
            
            # Add distance plot
            doc.add_heading(f"Distance Plot", level=2)
            doc.add_picture(traj['distance_plot_path'], width=Mm(150))
            
            # Add recurrence plot
            doc.add_heading(f"Recurrence Plot (Threshold: 0.3)", level=2)
            doc.add_picture(traj['recurrence_plot_path'], width=Mm(150))
            
            # Add page break after each trajectory except the last one
            if idx < len(context['trajectories']) - 1:
                doc.add_page_break()
        
        # Save the document
        doc_path = os.path.join(output_dir, f"launch_pentek_{TEMPERATURE}_{RADIUS_INITIAL_CONDITIONS}.docx")
        doc.save(doc_path)
        print(f"Created Word document using python-docx: {doc_path}")
        
    except Exception as e:
        print(f"Error with python-docx: {e}")
        
        # Alternative: Use a simpler approach with reportlab to create a PDF
        try:

            
            # Create PDF instead
            pdf_path = os.path.join(output_dir, f"launch_pentek_{TEMPERATURE}_{RADIUS_INITIAL_CONDITIONS}.pdf")
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title_style = styles['Title']
            story.append(Paragraph(context['title'], title_style))
            story.append(Spacer(1, 0.25*inch))
            
            # Add initial prompt if available
            if 'initial_prompt' in context:
                story.append(Paragraph("Initial Prompt", styles['Heading1']))
                story.append(Paragraph(context['initial_prompt'], styles['Normal']))
                story.append(Spacer(1, 0.5*inch))
            
            # Add trajectories
            for traj in context['trajectories']:
                idx = traj['index']
                story.append(Paragraph(f"Trajectory {idx}", styles['Heading1']))
                
                # Add generated text if available
                if 'generated_text' in traj:
                    story.append(Paragraph("Generated Text:", styles['Heading2']))
                    story.append(Paragraph(traj['generated_text'], styles['Normal']))
                    story.append(Spacer(1, 0.25*inch))
                
                # Add distance plot
                story.append(Paragraph(f"Distance Plot", styles['Heading2']))
                story.append(Image(traj['distance_plot_path'], width=6*inch, height=4.8*inch))
                story.append(Spacer(1, 0.25*inch))
                
                # Add recurrence plot
                story.append(Paragraph(f"Recurrence Plot (Threshold: 0.3)", styles['Heading2']))
                story.append(Image(traj['recurrence_plot_path'], width=6*inch, height=4.8*inch))
                
                # Add page break after each trajectory
                story.append(Spacer(1, 0.5*inch))
            
            # Build the PDF
            doc.build(story)
            print(f"Created PDF document instead: {pdf_path}")
            
        except Exception as e2:
            print(f"Error creating PDF: {e2}")
            print("Falling back to simple HTML export")
            
            # Last resort: create an HTML file
            html_path = os.path.join(output_dir, f"launch_pentek_{TEMPERATURE}_{RADIUS_INITIAL_CONDITIONS}.html")
            with open(html_path, 'w') as f:
                f.write(f"<html><head><title>{context['title']}</title></head><body>")
                f.write(f"<h1>{context['title']}</h1>")
                
                if 'initial_prompt' in context:
                    f.write("<h2>Initial Prompt</h2>")
                    f.write(f"<p>{context['initial_prompt']}</p>")
                    f.write("<hr>")
                
                for traj in context['trajectories']:
                    idx = traj['index']
                    f.write(f"<h2>Trajectory {idx}</h2>")
                    
                    if 'generated_text' in traj:
                        f.write("<h3>Generated Text:</h3>")
                        f.write(f"<p>{traj['generated_text']}</p>")
                    
                    f.write("<h3>Distance Plot</h3>")
                    f.write(f"<img src='{os.path.basename(traj['distance_plot_path'])}' width='800'>")
                    
                    f.write("<h3>Recurrence Plot (Threshold: 0.3)</h3>")
                    f.write(f"<img src='{os.path.basename(traj['recurrence_plot_path'])}' width='800'>")
                    
                    f.write("<hr>")
                
                f.write("</body></html>")
            print(f"Created HTML document as fallback: {html_path}")

# Main execution loop
for rrr in radiuses:
    for ttt in range(len(temps)):
        TEMPERATURE = temps[ttt]
        RADIUS_INITIAL_CONDITIONS = rrr
        RESULTS_DIR = f"./launch_pentek_{TEMPERATURE}_{RADIUS_INITIAL_CONDITIONS}/"
        PLOTS_DIR = f"./recurrence_plots_doc/launch_pentek_{TEMPERATURE}_{RADIUS_INITIAL_CONDITIONS}/"
        os.makedirs(PLOTS_DIR, exist_ok=True)
        
        # Load results file to get generated texts and initial prompt
        RESULTS_FILE = os.path.join(RESULTS_DIR, "results_summary.json")
        results = None
        try:
            with open(RESULTS_FILE, "r") as f:
                results = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"Warning: Could not load results file: {e}")
        
        # Load trajectories
        trajectories = load_pickle("trajectories")
        trajectories = trajectories[:11]  # Limit to 11 trajectories as in the original code
        
        # Generate plots and document
        recurrence_plot_with_threshold(trajectories, output_dir=PLOTS_DIR, results=results)